from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "ECG_Personal_Photo_Lock_Report.docx"

ACCENT = RGBColor(31, 111, 235)
DARK = RGBColor(36, 41, 47)
MUTED = RGBColor(87, 96, 106)
LIGHT_BLUE = "EAF2FF"
LIGHT_GREEN = "EAF7EF"
LIGHT_RED = "FCEEEE"
LIGHT_GRAY = "F6F8FA"


def font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_flowchart(filename: str, title: str, steps: list[str], colors: list[str] | None = None) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width = 1500
    height = 300
    margin = 40
    top = 88
    box_h = 105
    gap = 32
    n = len(steps)
    box_w = int((width - 2 * margin - gap * (n - 1)) / n)

    img = Image.new("RGB", (width, height), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font = font(34, True)
    box_font = font(23, True)
    small_font = font(18, False)

    d.text((margin, 28), title, fill=(36, 41, 47), font=title_font)
    d.line((margin, 72, width - margin, 72), fill=(208, 215, 222), width=2)

    if colors is None:
        colors = [LIGHT_BLUE] * n

    centers = []
    for i, step in enumerate(steps):
        x = margin + i * (box_w + gap)
        y = top
        fill = colors[i % len(colors)]
        outline = (31, 111, 235)
        d.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill="#" + fill, outline=outline, width=3)

        words = step.split()
        lines = []
        current = ""
        for word in words:
            test = (current + " " + word).strip()
            if d.textlength(test, font=box_font) <= box_w - 34:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        total_h = len(lines) * 28
        ty = y + (box_h - total_h) / 2 - 2
        for line in lines:
            tw = d.textlength(line, font=box_font)
            d.text((x + (box_w - tw) / 2, ty), line, fill=(36, 41, 47), font=box_font)
            ty += 30
        centers.append((x + box_w, y + box_h / 2, x + box_w + gap, y + box_h / 2))

    for i in range(n - 1):
        x1, y1, x2, y2 = centers[i]
        d.line((x1 + 5, y1, x2 - 15, y2), fill=(87, 96, 106), width=4)
        d.polygon([(x2 - 15, y2 - 10), (x2 - 15, y2 + 10), (x2 + 2, y2)], fill=(87, 96, 106))

    d.text((margin, height - 48), "Flowchart prepared for the ECG Personal Photo Lock report.", fill=(87, 96, 106), font=small_font)
    out = ASSET_DIR / filename
    img.save(out)
    return out


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, header_fill="EAF2FF"):
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = ACCENT if level == 1 else DARK
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return p


def add_placeholder_box(doc: Document, title: str, height_rows: int = 2):
    table = doc.add_table(rows=height_rows, cols=1)
    table.style = "Table Grid"
    for row in table.rows:
        row.height = Inches(0.45)
        cell = row.cells[0]
        set_cell_shading(cell, "FFF8E5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].cells[0].text = title
    for p in table.rows[0].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(10)
            r.font.color.rgb = MUTED
    doc.add_paragraph()


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = MUTED


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    flow_system = draw_flowchart(
        "flow_system.png",
        "System Overview",
        ["PTB Dataset", "Preprocessing", "Segmentation", "Wavelet Features", "Classifier", "Photo Unlock"],
        ["EAF2FF", "EAF7EF", "FFF8E5", "F2ECFF", "FCEEEE", "EAF7EF"],
    )
    flow_pre = draw_flowchart(
        "flow_preprocessing.png",
        "ECG Preprocessing Pipeline",
        ["Raw ECG", "Baseline Removal", "Bandpass Filter", "Notch Filter", "NeuroKit R-peaks", "Segments"],
        ["F6F8FA", "EAF7EF", "EAF2FF", "FFF8E5", "FCEEEE", "F2ECFF"],
    )
    flow_train = draw_flowchart(
        "flow_training.png",
        "Model Training and Comparison",
        ["Select Subjects", "Train/Test Split", "db1 db2 db4", "Train SVM KNN RF", "Compare Accuracy", "Best Model"],
        ["EAF2FF", "F6F8FA", "F2ECFF", "EAF7EF", "FFF8E5", "EAF7EF"],
    )
    flow_auth = draw_flowchart(
        "flow_authentication.png",
        "Authentication Decision Flow",
        ["Load ECG", "Segment Beats", "Predict Segments", "Majority Vote", "80% Rule", "Unlock or Reject"],
        ["F6F8FA", "EAF2FF", "F2ECFF", "FFF8E5", "EAF7EF", "FCEEEE"],
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"

    header = section.header.paragraphs[0]
    header.text = "ECG Personal Photo Lock - HCI Project Report"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Prepared for HCI Project - Idea 4"
    for r in footer.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ECG-Based Personal Photo Lock")
    run.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("HCI Project Report - Idea 4")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.cell(0, 0).text = "Project Name"
    meta.cell(0, 1).text = "ECG-Based Personal Photo Lock"
    meta.cell(1, 0).text = "Team Members"
    meta.cell(1, 1).text = "Name 1:\nName 2:\nName 3:\nName 4:"
    meta.cell(2, 0).text = "Dataset"
    meta.cell(2, 1).text = "PTB Diagnostic ECG Database, PhysioNet"
    meta.cell(3, 0).text = "Application Goal"
    meta.cell(3, 1).text = "Identify one of five trained subjects from ECG segments and display the matching photo."
    style_table(meta, "EAF2FF")

    add_heading(doc, "1. Project Overview", 1)
    add_body(doc, "This project implements an ECG-based personal photo lock. The system trains an identification model for five subjects, preprocesses ECG recordings, extracts wavelet features, compares several classifiers, and unlocks a photo panel only when a subject is confidently identified.")
    add_body(doc, "A subject is accepted when more than 80% of ECG heartbeat segments are classified as the same trained subject. If the segment agreement is below the threshold, or the selected ECG file belongs to a patient outside the trained set, the subject is treated as unidentified.")

    doc.add_picture(str(flow_system), width=Inches(6.6))
    add_caption(doc, "Figure 1. Overall ECG photo lock workflow.")

    add_heading(doc, "2. Data Preparation", 1)
    add_body(doc, "The data source is the PTB Diagnostic ECG Database from PhysioNet. Five subject folders are selected from the PTB patient records. In the current testing setup, the selected subjects are patient068, patient078, patient100, patient120, and patient140.")
    add_body(doc, "Each ECG record is read from WFDB format. Lead II is used for feature generation because it provides a clear ECG morphology and is commonly useful for R-peak detection. The extracted heartbeat feature rows are split into training and testing sets using a stratified split so each subject remains represented in both sets.")

    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Database", "PTB Diagnostic ECG Database"),
        ("Subjects", "Five PTB patients selected in the Training tab"),
        ("Lead", "Lead II, zero-based lead index 1"),
        ("Sampling Frequency", "1000 Hz"),
        ("Train/Test Split", "70% training, 30% testing"),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    style_table(table, "EAF7EF")

    add_heading(doc, "3. ECG Preprocessing", 1)
    add_body(doc, "The preprocessing stage reduces noise and prepares stable heartbeat windows for classification. The pipeline removes baseline wander, applies bandpass filtering, suppresses powerline noise, detects R-peaks using NeuroKit2, refines each marker to the dominant local ECG peak, and cuts fixed-length heartbeat segments around each detected R-peak.")
    doc.add_picture(str(flow_pre), width=Inches(6.6))
    add_caption(doc, "Figure 2. ECG preprocessing pipeline.")
    add_placeholder_box(doc, "Insert screenshot: raw ECG signal from ECG Viewer tab")
    add_placeholder_box(doc, "Insert screenshot: filtered ECG with detected R-peaks")
    add_placeholder_box(doc, "Insert screenshot: segmented heartbeat example or preprocessing result")

    add_heading(doc, "4. Feature Extraction", 1)
    add_body(doc, "The feature extraction method uses Discrete Wavelet Transform features. Three Daubechies mother wavelets are compared: db1, db2, and db4. For each heartbeat segment, wavelet coefficients are generated and summarized with statistical descriptors.")
    feat_table = doc.add_table(rows=6, cols=2)
    feat_rows = [
        ("Mother Wavelets", "db1, db2, db4"),
        ("Transform Type", "Discrete Wavelet Transform"),
        ("Decomposition Level", "Level 4"),
        ("Per-band Features", "Mean, standard deviation, energy, maximum absolute value, minimum absolute value"),
        ("Purpose", "Represent each ECG heartbeat segment numerically for classifier training"),
        ("Comparison", "The app trains and reports classifier performance separately for each wavelet"),
    ]
    for row, values in zip(feat_table.rows, feat_rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    style_table(feat_table, "F2ECFF")

    add_heading(doc, "5. Classifiers and Parameters", 1)
    add_body(doc, "Three classifiers are trained and compared. SVM is included as required, and KNN and Random Forest are used as additional classifiers. Several parameters are tested for each classifier to find the best result.")
    clf_table = doc.add_table(rows=4, cols=3)
    headers = ["Classifier", "Parameters Tested", "Role in Comparison"]
    for i, h in enumerate(headers):
        clf_table.rows[0].cells[i].text = h
    clf_rows = [
        ("SVM", "Kernel: RBF, linear, polynomial; C: 1, 10, 100; gamma: scale or auto", "Primary required classifier"),
        ("KNN", "Neighbors: 3, 5, 7, 11; metric: Euclidean or Manhattan", "Distance-based baseline"),
        ("Random Forest", "Trees: 50, 100, 200; max depth: None, 10, 20", "Tree ensemble baseline"),
    ]
    for r_i, values in enumerate(clf_rows, start=1):
        for c_i, value in enumerate(values):
            clf_table.rows[r_i].cells[c_i].text = value
    style_table(clf_table, "EAF2FF")

    doc.add_picture(str(flow_train), width=Inches(6.6))
    add_caption(doc, "Figure 3. Training and model comparison workflow.")

    add_heading(doc, "6. Classification Results", 1)
    add_body(doc, "The table below summarizes one verified training run using the selected five patients: patient068, patient078, patient100, patient120, and patient140. If the final demo is retrained with another set of patients, replace the values with the final Results tab output.")
    res_table = doc.add_table(rows=5, cols=5)
    for i, h in enumerate(["Classifier", "Best Wavelet", "Best Parameters", "Train Accuracy", "Test Accuracy"]):
        res_table.rows[0].cells[i].text = h
    result_rows = [
        ("SVM", "db1", "RBF kernel, C=1, gamma=scale", "100.00%", "100.00%"),
        ("KNN", "db1", "k=5, Manhattan metric", "99.89%", "100.00%"),
        ("Random Forest", "db2", "50 trees, max depth=None", "100.00%", "99.74%"),
        ("Overall Best", "db1", "SVM, RBF kernel, C=1, gamma=scale", "100.00%", "100.00%"),
    ]
    for r_i, values in enumerate(result_rows, start=1):
        for c_i, value in enumerate(values):
            res_table.rows[r_i].cells[c_i].text = value
    style_table(res_table, "EAF7EF")
    add_placeholder_box(doc, "Insert screenshot: Results tab comparison table", 3)

    add_heading(doc, "7. Authentication Interface", 1)
    add_body(doc, "The Lock Screen provides the authentication interface. After training, the user can scan a demo ECG signal or load an ECG file. The application preprocesses the signal, extracts wavelet features using the best model's wavelet, predicts each heartbeat segment, and applies the majority-vote rule.")
    doc.add_picture(str(flow_auth), width=Inches(6.6))
    add_caption(doc, "Figure 4. Authentication and unlock decision workflow.")
    add_body(doc, "When a known subject is identified, the matching subject photo is displayed. If the ECG belongs to a patient outside the trained five subjects, or if the confidence threshold is not met, access is denied and the subject is labeled unidentified.")
    add_placeholder_box(doc, "Insert screenshot: Lock Screen before scan")
    add_placeholder_box(doc, "Insert screenshot: successful unlock with subject photo")
    add_placeholder_box(doc, "Insert screenshot: unidentified / access denied case")
    add_placeholder_box(doc, "Insert screenshot: Subject Manager with assigned photos")

    add_heading(doc, "8. Requirement Coverage", 1)
    cov = doc.add_table(rows=8, cols=3)
    for i, h in enumerate(["No.", "Requirement", "Implemented In Project"]):
        cov.rows[0].cells[i].text = h
    cov_rows = [
        ("1", "Identification interface for five subjects", "Training tab selects five PTB patients and Lock Screen performs identification"),
        ("2", "Suitable ECG preprocessing", "Baseline removal, bandpass filter, notch filter, NeuroKit R-peak detection, local peak refinement, segmentation"),
        ("3", "Wavelet features comparing db1, db2, db4", "Feature extraction and Results tab compare all three wavelets"),
        ("4", "Three classifiers including SVM with parameter tests", "SVM, KNN, and Random Forest parameter grids are trained"),
        ("5", "More than 80% segment agreement for identification", "Majority-vote rule in authentication logic"),
        ("6", "Compare classifiers by accuracy", "Results tab and report table compare train/test accuracy"),
        ("7", "Display photo when subject is present", "Subject Manager assigns photos; Lock Screen displays identified subject photo"),
    ]
    for r_i, values in enumerate(cov_rows, start=1):
        for c_i, value in enumerate(values):
            cov.rows[r_i].cells[c_i].text = value
    style_table(cov, "FFF8E5")

    add_heading(doc, "9. Conclusion", 1)
    add_body(doc, "The ECG Personal Photo Lock demonstrates a biometric authentication interface based on ECG signals. The system integrates ECG preprocessing, wavelet-based feature extraction, classifier comparison, confidence-based identification, unknown-subject rejection, and photo presentation for identified subjects.")

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
